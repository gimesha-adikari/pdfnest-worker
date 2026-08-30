import os
import tempfile
import fitz
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from unittest.mock import patch
from concurrent.futures import ThreadPoolExecutor

from app.api.tools.pdf_to_office.converters.word import (
    _get_pdf2docx_worker_count,
    _requires_structured_ocr,
    convert_to_word,
)


def create_sample_text_pdf(path: str, pages: int = 5) -> None:
    doc = fitz.open()
    for i in range(pages):
        page = doc.new_page()
        page.insert_text((50, 50), f"Page {i+1} Sample Text Content for PDF to Word Parallelism Unit Test.\n" * 5)
    doc.save(path)
    doc.close()


def create_scanned_text_pdf(path: str, text: str = "OCR V2 scanned Word text") -> None:
    image = Image.new("RGB", (1800, 500), "white")
    draw = ImageDraw.Draw(image)
    font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 64)
    draw.text((80, 180), text, fill="black", font=font)

    doc = fitz.open()
    page = doc.new_page(width=900, height=250)
    page.insert_image(page.rect, stream=_png_bytes(image))
    doc.save(path)
    doc.close()


def _png_bytes(image: Image.Image) -> bytes:
    from io import BytesIO

    buffer = BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


def create_mixed_text_pdf(path: str) -> None:
    doc = fitz.open()
    page = doc.new_page(width=600, height=300)
    page.insert_text((50, 100), "Native page content")

    image = Image.new("RGB", (1200, 400), "white")
    draw = ImageDraw.Draw(image)
    font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 56)
    draw.text((50, 150), "Scanned page content", fill="black", font=font)
    page = doc.new_page(width=600, height=200)
    page.insert_image(page.rect, stream=_png_bytes(image))
    doc.save(path)
    doc.close()


def test_get_pdf2docx_worker_count():
    with patch("os.cpu_count", return_value=1):
        assert _get_pdf2docx_worker_count() == 1

    with patch("os.cpu_count", return_value=2):
        assert _get_pdf2docx_worker_count() == 2

    with patch("os.cpu_count", return_value=8):
        assert _get_pdf2docx_worker_count() == 2

    with patch("os.cpu_count", return_value=None):
        assert _get_pdf2docx_worker_count() == 1


def test_convert_to_word_isolated_execution():
    with tempfile.TemporaryDirectory() as tmpdir:
        pdf_path = os.path.join(tmpdir, "test.pdf")
        docx_path = os.path.join(tmpdir, "test.docx")
        create_sample_text_pdf(pdf_path, pages=3)

        initial_files = set(os.listdir(os.getcwd()))
        convert_to_word(pdf_path, docx_path)

        assert os.path.exists(docx_path)
        assert os.path.getsize(docx_path) > 0
        final_files = set(os.listdir(os.getcwd()))
        # Verify no intermediate pages-*.json files leaked into CWD
        assert final_files == initial_files


def test_concurrent_pdf_to_word_no_cwd_collision():
    with tempfile.TemporaryDirectory() as tmpdir:
        pdf1 = os.path.join(tmpdir, "doc1.pdf")
        pdf2 = os.path.join(tmpdir, "doc2.pdf")
        docx1 = os.path.join(tmpdir, "out1.docx")
        docx2 = os.path.join(tmpdir, "out2.docx")

        create_sample_text_pdf(pdf1, pages=5)
        create_sample_text_pdf(pdf2, pages=5)

        with ThreadPoolExecutor(max_workers=2) as executor:
            f1 = executor.submit(convert_to_word, pdf1, docx1)
            f2 = executor.submit(convert_to_word, pdf2, docx2)
            f1.result()
            f2.result()

        assert os.path.exists(docx1) and os.path.getsize(docx1) > 0
        assert os.path.exists(docx2) and os.path.getsize(docx2) > 0


def test_scanned_pdf_to_word_uses_structured_ocr_v2():
    with tempfile.TemporaryDirectory() as tmpdir:
        pdf_path = os.path.join(tmpdir, "scanned.pdf")
        docx_path = os.path.join(tmpdir, "scanned.docx")
        create_scanned_text_pdf(pdf_path)

        assert _requires_structured_ocr(pdf_path) is True
        convert_to_word(pdf_path, docx_path, language="eng")

        output = Document(docx_path)
        text = "\n".join(paragraph.text for paragraph in output.paragraphs)
        assert "OCR V2" in text
        assert "scanned Word text" in text
        assert output.tables == []


def test_mixed_pdf_to_word_preserves_native_and_ocr_pages_in_order():
    with tempfile.TemporaryDirectory() as tmpdir:
        pdf_path = os.path.join(tmpdir, "mixed.pdf")
        docx_path = os.path.join(tmpdir, "mixed.docx")
        create_mixed_text_pdf(pdf_path)

        convert_to_word(pdf_path, docx_path, language="eng")

        output = Document(docx_path)
        text = "\n".join(paragraph.text for paragraph in output.paragraphs)
        assert text.index("Native page content") < text.index("Scanned page content")

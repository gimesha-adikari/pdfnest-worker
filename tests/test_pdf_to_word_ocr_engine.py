from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace

import pytest

import app.core.pdf_to_word_ocr_engine as engine


def test_pdf_to_word_ocr_engine_defaults_to_internal(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.delenv(engine.PDF_TO_WORD_OCR_ENGINE_ENV, raising=False)

    assert engine.configured_pdf_to_word_ocr_engine() == "internal"


def test_pdf_to_word_ocr_engine_accepts_sdk(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv(engine.PDF_TO_WORD_OCR_ENGINE_ENV, " SDK ")

    assert engine.configured_pdf_to_word_ocr_engine() == "sdk"


def test_pdf_to_word_ocr_engine_rejects_unknown_value(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv(engine.PDF_TO_WORD_OCR_ENGINE_ENV, "external")

    with pytest.raises(engine.PdfToWordOcrEngineConfigurationError, match="PDF_TO_WORD_OCR_ENGINE"):
        engine.configured_pdf_to_word_ocr_engine()


def test_internal_selector_uses_frozen_structured_processor(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    calls: dict[str, object] = {}
    marker = object()

    class FakeProcessor:
        def process_document(self, path: str | Path, **kwargs: object) -> object:
            calls["path"] = path
            calls.update(kwargs)
            return marker

    monkeypatch.setenv(engine.PDF_TO_WORD_OCR_ENGINE_ENV, "internal")
    monkeypatch.setattr(engine, "_internal_processor", lambda: FakeProcessor())

    result = engine.execute_pdf_to_word_ocr(tmp_path / "source.pdf", language="eng")

    assert result is marker
    assert calls == {"path": tmp_path / "source.pdf", "language": "eng"}


def test_sdk_selector_uses_public_processor_and_not_internal(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    calls: dict[str, object] = {}
    marker = object()

    class FakeProcessor:
        def extract_document(self, path: str | Path, **kwargs: object) -> object:
            calls["path"] = path
            calls.update(kwargs)
            return marker

    monkeypatch.setenv(engine.PDF_TO_WORD_OCR_ENGINE_ENV, "sdk")
    monkeypatch.setattr(engine, "_sdk_processor", lambda: FakeProcessor())
    monkeypatch.setattr(
        engine,
        "_internal_processor",
        lambda: pytest.fail("PDF-to-Word SDK mode must not construct the internal processor"),
    )

    result = engine.execute_pdf_to_word_ocr(tmp_path / "source.pdf", language="eng+sin")

    assert result is marker
    assert calls == {"path": tmp_path / "source.pdf", "language": "eng+sin"}


def test_sdk_failure_does_not_fall_back_to_internal(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    class FakeProcessor:
        def extract_document(self, _path: str | Path, **_kwargs: object) -> object:
            raise RuntimeError("selected SDK failed")

    monkeypatch.setenv(engine.PDF_TO_WORD_OCR_ENGINE_ENV, "sdk")
    monkeypatch.setattr(engine, "_sdk_processor", lambda: FakeProcessor())
    monkeypatch.setattr(
        engine,
        "_internal_processor",
        lambda: pytest.fail("SDK failure must not fall back to internal"),
    )

    with pytest.raises(engine.PdfToWordOcrEngineExecutionError, match="PDF-to-Word OCR processing failed"):
        engine.execute_pdf_to_word_ocr(tmp_path / "source.pdf")


def test_sdk_failure_projection_does_not_expose_raw_exception(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    class FakeProcessor:
        def extract_document(self, _path: str | Path, **_kwargs: object) -> object:
            raise RuntimeError("/private/path and platen_document internals")

    monkeypatch.setenv(engine.PDF_TO_WORD_OCR_ENGINE_ENV, "sdk")
    monkeypatch.setattr(engine, "_sdk_processor", lambda: FakeProcessor())

    with pytest.raises(engine.PdfToWordOcrEngineExecutionError) as error:
        engine.execute_pdf_to_word_ocr(tmp_path / "source.pdf")

    assert str(error.value) == "PDF-to-Word OCR processing failed"


def test_docx_projection_accepts_sdk_canonical_element_values(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    from docx import Document
    from app.api.tools.pdf_to_office.converters.word import _convert_structured_to_word

    class FakeProcessor:
        def extract_document(self, _path: str | Path, **_kwargs: object) -> object:
            heading = SimpleNamespace(
                element_id="heading",
                type=SimpleNamespace(value="HEADING"),
                text="SDK heading",
                level=2,
                data={},
                ordered=None,
            )
            paragraph = SimpleNamespace(
                element_id="paragraph",
                type=SimpleNamespace(value="PARAGRAPH"),
                text="SDK paragraph",
                level=None,
                data={},
                ordered=None,
            )
            page = SimpleNamespace(
                elements=(heading, paragraph),
                reading_order=("heading", "paragraph"),
            )
            return SimpleNamespace(pages=(page,))

    monkeypatch.setenv(engine.PDF_TO_WORD_OCR_ENGINE_ENV, "sdk")
    monkeypatch.setattr(engine, "_sdk_processor", lambda: FakeProcessor())

    output_path = tmp_path / "result.docx"
    _convert_structured_to_word(str(tmp_path / "source.pdf"), str(output_path), "eng")

    output = Document(output_path)
    assert [paragraph.text for paragraph in output.paragraphs] == ["SDK heading", "SDK paragraph"]
    assert output.paragraphs[0].style.name.startswith("Heading")


def test_pdf_to_word_structured_projection_extracts_once(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    import app.api.tools.pdf_to_office.converters.word as word

    calls: list[tuple[object, object]] = []
    element = SimpleNamespace(
        element_id="paragraph",
        type=SimpleNamespace(value="PARAGRAPH"),
        text="one-pass result",
        level=None,
        data={},
        ordered=None,
    )
    result = SimpleNamespace(
        pages=(
            SimpleNamespace(
                elements=(element,),
                reading_order=("paragraph",),
            ),
        ),
    )

    def execute(path: str | Path, *, language: str) -> object:
        calls.append((path, language))
        return result

    monkeypatch.setattr(word, "execute_pdf_to_word_ocr", execute)
    output_path = tmp_path / "result.docx"
    word._convert_structured_to_word(str(tmp_path / "source.pdf"), str(output_path), "eng")

    assert calls == [(str(tmp_path / "source.pdf"), "eng")]
    assert output_path.exists()

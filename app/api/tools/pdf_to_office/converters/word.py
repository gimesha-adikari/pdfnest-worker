from __future__ import annotations

import os
import sys
import tempfile
import subprocess
from typing import Any
import fitz
from docx import Document

from app.core.ocr_v2.native import NativeDecision, NativeExtractor, NativeValidator
from app.core.pdf_to_word_ocr_engine import (
    configured_pdf_to_word_ocr_engine,
    execute_pdf_to_word_ocr,
)


def _get_pdf2docx_worker_count() -> int:
    cpu_count = os.cpu_count() or 1
    return min(2, max(1, cpu_count))


def _run_pdf2docx_isolated(pdf_path: str, output_path: str, workers: int) -> None:
    abs_pdf = os.path.abspath(pdf_path)
    abs_output = os.path.abspath(output_path)

    use_mp = workers > 1

    with tempfile.TemporaryDirectory(prefix="pdf2docx-job-") as job_dir:
        py_script = (
            "from pdf2docx import Converter\n"
            f"cv = Converter({abs_pdf!r})\n"
            "try:\n"
            "    kwargs = {\n"
            "        'keep_page_layout': False,\n"
            "        'connected_border': True,\n"
            "        'line_overlap_margin': 0.2,\n"
            "        'line_margin': 0.2,\n"
            "        'word_margin': 0.2,\n"
            "        'bottom_margin': 5.0,\n"
            f"        'multi_processing': {use_mp!r},\n"
            f"        'cpu_count': {workers!r},\n"
            "    }\n"
            f"    cv.convert({abs_output!r}, start=0, end=None, **kwargs)\n"
            "finally:\n"
            "    cv.close()\n"
        )
        proc = subprocess.Popen(
            [sys.executable, "-c", py_script],
            cwd=job_dir,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
        _, stderr = proc.communicate()
        if proc.returncode != 0:
            err_msg = stderr.decode("utf-8", errors="replace").strip()
            raise RuntimeError(f"PDF to Word conversion failed ({proc.returncode}): {err_msg}")


def _requires_structured_ocr(pdf_path: str) -> bool:
    """Select the structured OCR route only for scanned or mixed pages.

    Classification is native-only and therefore does not invoke OCR.  The
    structured processor then performs the one native-first OCR V2 pass for
    pages that actually need it.
    """
    extractor = NativeExtractor()
    validator = NativeValidator()
    with fitz.open(pdf_path) as doc:
        for page_index, page in enumerate(doc):
            candidate = extractor.extract(page, page_index)
            decision = validator.validate(candidate)
            if decision.decision != NativeDecision.TRUST_NATIVE:
                return True
    return False


def _structured_element_type(element: Any) -> str:
    """Read the canonical element value across internal and SDK enum types."""

    return str(getattr(getattr(element, "type", None), "value", getattr(element, "type", ""))).upper()


def _add_structured_element(doc: Document, element: Any) -> None:
    """Map only structure represented by the canonical structured result."""
    element_type = _structured_element_type(element)
    if element_type == "HEADING":
        doc.add_heading(element.text, level=max(1, min(9, element.level or 1)))
    elif element_type in {"PARAGRAPH", "TEXT_BLOCK"}:
        if element.text:
            doc.add_paragraph(element.text)
    elif element_type == "LIST":
        items = element.data.get("items", [])
        style = "List Number" if element.ordered else "List Bullet"
        for item in items:
            text = str(item.get("text", "")).strip()
            if text:
                doc.add_paragraph(text, style=style)
    elif element_type == "TABLE":
        headers = element.data.get("headers", [])
        rows = element.data.get("rows", [])
        table_rows = ([headers] if headers else []) + list(rows)
        column_count = max((len(row) for row in table_rows), default=0)
        if column_count:
            table = doc.add_table(rows=0, cols=column_count)
            table.style = "Table Grid"
            for row in table_rows:
                cells = table.add_row().cells
                for index, cell in enumerate(row):
                    cells[index].text = str(cell.get("text", ""))
    elif element_type == "CAPTION":
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(element.text)
        run.italic = True
    elif element_type == "FORMULA":
        # Only genuine structured formula text reaches this mapper.  Current
        # local Tesseract structured output deliberately does not fabricate it.
        if element.text:
            doc.add_paragraph(element.text)


def _write_structured_result_to_word(result: Any, output_path: str) -> None:
    doc_out = Document()
    for page_index, page in enumerate(result.pages):
        elements = {element.element_id: element for element in page.elements}
        for element_id in page.reading_order:
            element = elements[element_id]
            _add_structured_element(doc_out, element)
        if page_index < len(result.pages) - 1:
            doc_out.add_page_break()
    doc_out.save(output_path)


def _convert_structured_to_word(pdf_path: str, output_path: str, language: str) -> None:
    result = execute_pdf_to_word_ocr(pdf_path, language=language)
    _write_structured_result_to_word(result, output_path)


def convert_to_word(pdf_path: str, output_path: str, language: str = "eng") -> None:
    # Validate the consumer selector even when this document takes the native
    # route, so an invalid deployment configuration cannot be hidden by the
    # absence of an OCR fallback on a particular input.
    configured_pdf_to_word_ocr_engine()
    doc = fitz.open(pdf_path)
    try:
        structured = _requires_structured_ocr(pdf_path)
    finally:
        doc.close()

    if structured:
        _convert_structured_to_word(pdf_path, output_path, language)
        return

    workers = _get_pdf2docx_worker_count()
    _run_pdf2docx_isolated(pdf_path, output_path, workers)
